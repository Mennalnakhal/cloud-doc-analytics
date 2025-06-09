from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
from PyPDF2 import PdfReader
from docx import Document
from flask_cors import CORS


app = Flask(__name__)
CORS(app)
UPLOAD_FOLDER = 'data'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def extract_title_from_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        if reader.pages:
            first_page = reader.pages[0]
            text = first_page.extract_text()
            if text:
                lines = text.strip().split('\n')
                return lines[0] if lines else "No title found"
    except Exception as e:
        return f"PDF error: {str(e)}"
    return "No title found"

def extract_title_from_docx(filepath):
    try:
        doc = Document(filepath)
        if doc.paragraphs:
            return doc.paragraphs[0].text.strip()
    except Exception as e:
        return f"DOCX error: {str(e)}"
    return "No title found"

@app.route('/upload', methods=['POST'])
def upload():
    file = request.files['file']
    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    if filename.lower().endswith('.pdf'):
        title = extract_title_from_pdf(filepath)
    elif filename.lower().endswith('.docx'):
        title = extract_title_from_docx(filepath)
    else:
        return jsonify({"error": "Unsupported file type"}), 400

    return jsonify({"filename": filename, "title": title})

if __name__ == '__main__':
    app.run(debug=True)
